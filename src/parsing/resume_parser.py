"""
resume_parser.py — Extract plain text from PDF, DOCX, or .txt resume files.

Usage:
    from src.parsing.resume_parser import parse_resume
    text = parse_resume("/path/to/resume.pdf")
"""

import io
import pathlib


# ── PDF ───────────────────────────────────────────────────────────────────────

def _parse_pdf(file_path: str | pathlib.Path | bytes | io.BytesIO) -> str:
    """
    Extract text from a PDF.
    Tries pdfminer.six first (higher quality), falls back to PyPDF2.
    """
    # Accept raw bytes / BytesIO (for Streamlit file uploader)
    if isinstance(file_path, (bytes, bytearray)):
        file_like = io.BytesIO(file_path)
    elif isinstance(file_path, io.BytesIO):
        file_like = file_path
    else:
        file_like = open(file_path, "rb")

    text = ""

    # Try pdfminer.six
    try:
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams
        output = io.StringIO()
        file_like.seek(0)
        extract_text_to_fp(file_like, output, laparams=LAParams())
        text = output.getvalue()
    except Exception:
        pass

    # Fallback to PyPDF2
    if not text.strip():
        try:
            import PyPDF2
            file_like.seek(0)
            reader = PyPDF2.PdfReader(file_like)
            text   = "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        except Exception:
            pass

    # Fallback to pypdf
    if not text.strip():
        try:
            import pypdf
            file_like.seek(0)
            reader = pypdf.PdfReader(file_like)
            text   = "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        except Exception:
            pass

    if not isinstance(file_path, (bytes, bytearray, io.BytesIO)):
        file_like.close()

    return text.strip()


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _parse_docx(file_path: str | pathlib.Path | bytes | io.BytesIO) -> str:
    """Extract text from a Word (.docx) file."""
    try:
        from docx import Document
        if isinstance(file_path, (bytes, bytearray)):
            doc = Document(io.BytesIO(file_path))
        elif isinstance(file_path, io.BytesIO):
            doc = Document(file_path)
        else:
            doc = Document(str(file_path))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as exc:
        raise RuntimeError(f"Failed to parse DOCX: {exc}") from exc


# ── TXT ───────────────────────────────────────────────────────────────────────

def _parse_txt(file_path: str | pathlib.Path | bytes | io.BytesIO) -> str:
    if isinstance(file_path, (bytes, bytearray)):
        return file_path.decode("utf-8", errors="ignore")
    if isinstance(file_path, io.BytesIO):
        return file_path.read().decode("utf-8", errors="ignore")
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


# ── Public API ────────────────────────────────────────────────────────────────

def parse_resume(
    file: str | pathlib.Path | bytes | io.BytesIO,
    file_type: str | None = None,
) -> str:
    """
    Extract plain text from a resume file.

    Parameters
    ----------
    file      : file path (str/Path), raw bytes, or BytesIO object
    file_type : 'pdf', 'docx', or 'txt'.
                If None, inferred from file extension (path) or content.

    Returns
    -------
    Extracted text string.

    Raises
    ------
    ValueError if the file type cannot be determined or is unsupported.
    """
    # Infer type from path
    if file_type is None:
        if isinstance(file, (str, pathlib.Path)):
            suffix = pathlib.Path(file).suffix.lower().lstrip(".")
            file_type = suffix
        else:
            # Can't infer from bytes — default to pdf
            file_type = "pdf"

    file_type = file_type.lower().lstrip(".")

    if file_type == "pdf":
        text = _parse_pdf(file)
    elif file_type in ("docx", "doc"):
        text = _parse_docx(file)
    elif file_type == "txt":
        text = _parse_txt(file)
    else:
        raise ValueError(
            f"Unsupported file type: '{file_type}'. "
            "Only pdf, docx, and txt are supported."
        )

    if not text.strip():
        raise ValueError("Could not extract any text from the uploaded file. "
                         "Ensure the document is not a scanned image.")

    return text
